from textblob import TextBlob
import pandas as pd
import streamlit as st
from docx import Document
import PyPDF2
import re
from bs4 import BeautifulSoup
import requests
from PIL import Image
import easyocr
import numpy as np
from tensorflow.keras.models import load_model
from tensorflow.keras.preprocessing.sequence import pad_sequences
from tensorflow.keras.preprocessing.text import Tokenizer

# Initialize EasyOCR reader (supports English by default, add languages as needed)
reader = easyocr.Reader(['en'])  # You can add more languages, e.g., ['en', 'fr', 'de']

# Load the pre-trained sentiment analysis model (.h5)
model = load_model(r'./sentiment_lstm_model.h5')

# Load tokenizer 
tokenizer = Tokenizer()

# Sentiment analysis function
def predict_sentiment(text):
    # Preprocess the input text (same steps as during training)
    sequences = tokenizer.texts_to_sequences([text])
    padded_sequences = pad_sequences(sequences, maxlen=1000)  # Adjust maxlen based on your model's input size
    
    # Predict sentiment using the loaded model
    prediction = model.predict(padded_sequences)
    
    # For multi-class classification (3 classes: Negative, Neutral, Positive)
    sentiment_classes = ["Negative", "Neutral", "Positive"]
    predicted_class_index = np.argmax(prediction[0])  # Get the index of the highest probability
    sentiment = sentiment_classes[predicted_class_index]  # Get the corresponding sentiment label
    
    return sentiment, prediction[0]  # Return sentiment and prediction probabilities


# File uploader for multiple files
st.header('📝 Sentiment Analysis for files upload')
uploaded_files = st.file_uploader(
    "Upload multiple files (CSV, Excel, TXT, DOCX, PDF, Images - JPG, PNG)",
    type=["csv", "xlsx", "xls", "txt", "docx", "pdf", "jpg", "jpeg", "png"],
    accept_multiple_files=True
)

if uploaded_files:
    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name.lower()
        st.subheader(f"📄 Processing: {uploaded_file.name}")

        try:
            # CSV
            # will read row by row
            if file_name.endswith(".csv"):
                df = pd.read_csv(uploaded_file)
                st.write("✅ CSV file loaded successfully!")
            
            # Excel
            # will read row by row
            elif file_name.endswith((".xlsx", ".xls")):
                df = pd.read_excel(uploaded_file)
                st.write("✅ Excel file loaded successfully!")
            
            # Text
            # will read line by line (as long as the paragraph are together)
            elif file_name.endswith(".txt"):
                content = uploaded_file.read().decode("utf-8")
                df = pd.DataFrame({"text": content.splitlines()})
                st.write("✅ Text file loaded successfully!")
            
            # Word Document
            # look for each paragraph in doc.paragraphs
            # p.text extracts the text from each paragraph while empty paragraphs (only spaces/newlines) are removed
            # then takes all the extracted text paragraph and joins them into a single string, with each paragraph separated by a newline
            elif file_name.endswith(".docx"):
                doc = Document(uploaded_file)
                content = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
                df = pd.DataFrame({"text": content.splitlines()})
                st.write("✅ Word document loaded successfully!")
            
            # PDF
            #need split paragraphs by double newlines to analyze for each paragraph. Otherwise, it will take as one whole paragraph
            elif file_name.endswith(".pdf"):
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                pdf_text = ""

                # Extract text from each page
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        pdf_text += text + "\n"  # Add newline to maintain structure

                # Split paragraphs by at least double newlines
                paragraphs = re.split(r'\n\s*\n', pdf_text.strip())  # handle inconsistent spacing between paragraphs

                # Convert to DataFrame
                df = pd.DataFrame({"text": paragraphs})
                
                st.write("✅ PDF file loaded successfully!")
            
            #image files(JPG,PNG,JPEG) - extract text using OCR
            elif file_name.endswith((".jpg", ".jpeg", ".png")):
                image = Image.open(uploaded_file)
                st.image(image, caption="Uploaded Image", use_container_width=True)
               
                # Convert PIL Image to NumPy Array for EasyOCR
                image_np = np.array(image)
          
                # Extract text using EasyOCR
                extracted_text = reader.readtext(image_np, detail=0)  # `detail=0` returns just the text
                
                
                if extracted_text:
                    df = pd.DataFrame({"text": extracted_text})
                    st.success("✅ Image text extracted successfully!")
                else:
                    st.error("❌ No readable text found in the image!")
                    continue
                
            else:
                st.error(f"❌ Unsupported file type for {uploaded_file.name}")
                continue

            # Check for available text columns
            # for files without any text
            text_columns = [col for col in df.columns if df[col].dtype == 'object']
            if not text_columns:
                st.error(f"❌ The file {uploaded_file.name} does not contain any valid text columns for analysis.")
                continue

            # Use the first available text column for sentiment analysis, convert to string type
            text_column = text_columns[0]
            df[text_column] = df[text_column].astype(str)

            # Apply sentiment analysis
            sentiments = []
            probabilities = []
            for text in df[text_column]:
                sentiment, probability = predict_sentiment(text)
                sentiments.append(sentiment)
                #probabilities.append(probability)
            
            # Add sentiment and probabilities to the DataFrame
            df['analysis'] = sentiments
            #df['probability'] = probabilities
            
            # Add a column with today's date
            # Add today's date to each row
            #normalize for time to become 00:00:00
            # convert to readable format
            df['date'] = pd.to_datetime('today').normalize().strftime('%Y-%m-%d %H:%M:%S') 
            
            # Show text with respective score and analysis
            st.write("Analyzed Data:")
            df_result = df[[text_column, 'analysis', 'date']]
            st.dataframe(df_result.head(1000)) #take first 1000 rows

            # Download the processed file
            @st.cache_data
            def convert_df_to_csv(df):
                return df.to_csv(index=False).encode('utf-8')

            #maximum six can be download
            csv_data = convert_df_to_csv(df_result)
            st.download_button(
                label=f"📥 Download results for {uploaded_file.name}",
                data=csv_data,
                file_name=f"{uploaded_file.name.split('.')[0]}_sentiment.csv",
                mime="text/csv"
            )

        except Exception as e:
            st.error(f"⚠️ Error processing {uploaded_file.name}: {e}")
